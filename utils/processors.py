import os, pandas as pd, base64
import pdfplumber
from docx import Document
import re
from typing import List, Dict, Any

# OCR imports (optional - for scanned/image-based PDFs)
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️  OCR not available. Install pytesseract and pdf2image for scanned PDF support.")

DOWNLOADS_DIR = "downloads"
os.makedirs(DOWNLOADS_DIR, exist_ok=True)

def save_attachment(attachment):
    """Save email attachment to downloads directory"""
    filename = attachment["name"]
    content_bytes = attachment.get("contentBytes")
    if not content_bytes:
        return None
    
    try:
        # Handle base64 encoded content
        if isinstance(content_bytes, str):
            # If it's base64 encoded, decode it
            try:
                decoded_content = base64.b64decode(content_bytes)
            except:
                # If decoding fails, treat as regular string
                decoded_content = content_bytes.encode('utf-8')
        else:
            decoded_content = content_bytes
            
        path = os.path.join(DOWNLOADS_DIR, filename)
        with open(path, "wb") as f:
            f.write(decoded_content)
        
        print(f"✅ Saved attachment: {filename}")
        return path
    except Exception as e:
        print(f"❌ Error saving attachment {filename}: {e}")
        return None

def _format_table_as_text(table: list) -> str:
    """Convert a table (list of rows) to pipe-delimited text format"""
    if not table:
        return ""

    formatted_rows = []
    for row in table:
        # Clean each cell: replace None with empty string, strip whitespace
        cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
        formatted_rows.append(" | ".join(cleaned_row))

    return "\n".join(formatted_rows)


def _extract_text_with_ocr(path: str) -> str:
    """
    Extract text from image-based/scanned PDF using OCR.
    Requires: pytesseract, pdf2image, and Tesseract OCR installed.
    """
    if not OCR_AVAILABLE:
        print("⚠️  OCR libraries not installed. Cannot process scanned PDF.")
        return ""

    try:
        print("🔍 Attempting OCR extraction for scanned PDF...")
        extracted_text = []

        # Convert PDF pages to images
        images = convert_from_path(path, dpi=300)

        for page_num, image in enumerate(images, 1):
            page_content = [f"=== PAGE {page_num} (OCR) ==="]

            # Run OCR on the image
            page_text = pytesseract.image_to_string(image, lang='eng')

            if page_text.strip():
                page_content.append(page_text)
                extracted_text.append("\n".join(page_content))

        full_text = "\n\n".join(extracted_text)
        print(f"✅ Extracted text from PDF using OCR: {len(full_text)} characters")
        return full_text

    except Exception as e:
        print(f"❌ OCR extraction failed: {e}")
        return ""


def extract_text_from_pdf(path: str) -> str:
    """
    Extract text from PDF files with enhanced table extraction using pdfplumber.
    Falls back to OCR for scanned/image-based PDFs.

    This function:
    1. Attempts to detect and extract tables from each page using pdfplumber
    2. Extracts remaining non-table text
    3. If no text extracted (scanned PDF), falls back to OCR
    4. Combines all into a structured format that's easier for AI to parse
    """
    try:
        extracted_text = []

        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_content = []
                page_content.append(f"=== PAGE {page_num + 1} ===")

                # Try to extract tables from this page
                tables = page.extract_tables(table_settings={
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "snap_tolerance": 3,
                    "join_tolerance": 3,
                    "edge_min_length": 3,
                    "min_words_vertical": 1,
                    "min_words_horizontal": 1,
                })

                # If no tables found with lines strategy, try text-based detection
                if not tables:
                    tables = page.extract_tables(table_settings={
                        "vertical_strategy": "text",
                        "horizontal_strategy": "text",
                        "snap_tolerance": 3,
                        "join_tolerance": 3,
                    })

                if tables:
                    # Extract and format tables
                    for table_idx, table in enumerate(tables, 1):
                        if table and len(table) > 0:
                            table_text = _format_table_as_text(table)
                            if table_text.strip():
                                page_content.append(f"=== TABLE {table_idx} ===")
                                page_content.append(table_text)

                # Also extract regular text (for content outside tables)
                page_text = page.extract_text() or ""
                if page_text.strip():
                    # If we have tables, label the text section
                    if tables:
                        page_content.append("=== TEXT CONTENT ===")
                    page_content.append(page_text)

                if len(page_content) > 1:  # More than just the page header
                    extracted_text.append("\n".join(page_content))

        full_text = "\n\n".join(extracted_text)
        print(f"✅ Extracted text from PDF using pdfplumber: {len(full_text)} characters")

        # If pdfplumber extracted no text, try OCR (for scanned/image PDFs)
        if len(full_text.strip()) == 0:
            print("⚠️  No text extracted with pdfplumber. PDF may be scanned/image-based.")
            full_text = _extract_text_with_ocr(path)

        return full_text

    except Exception as e:
        print(f"❌ Error extracting PDF text from {path}: {e}")
        # Try OCR as last resort
        return _extract_text_with_ocr(path)

def extract_text_from_excel(path: str) -> str:
    """Extract text from Excel files, handling multiple sheets"""
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(path)
        extracted_data = []
        
        for sheet_name in excel_file.sheet_names:
            try:
                df = pd.read_excel(path, sheet_name=sheet_name)
                if not df.empty:
                    # Clean up the dataframe
                    df = df.fillna('')  # Replace NaN with empty strings
                    
                    # Convert to structured text
                    sheet_text = f"=== SHEET: {sheet_name} ===\n"
                    sheet_text += df.to_string(index=False)
                    extracted_data.append(sheet_text)
            except Exception as e:
                print(f"⚠️  Warning: Could not read sheet '{sheet_name}': {e}")
                continue
        
        full_text = "\n\n".join(extracted_data)
        print(f"✅ Extracted text from Excel: {len(full_text)} characters")
        return full_text
    except Exception as e:
        print(f"❌ Error extracting Excel text from {path}: {e}")
        return ""

def extract_text_from_docx(path: str) -> str:
    """Extract text from Word documents"""
    try:
        doc = Document(path)
        extracted_text = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                extracted_text.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            table_text = "=== TABLE ===\n"
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                table_text += row_text + "\n"
            extracted_text.append(table_text)
        
        full_text = "\n".join(extracted_text)
        print(f"✅ Extracted text from Word doc: {len(full_text)} characters")
        return full_text
    except Exception as e:
        print(f"❌ Error extracting Word doc text from {path}: {e}")
        return ""

def extract_text_from_txt(path: str) -> str:
    """Extract text from TXT files with encoding detection"""
    try:
        # Try different encodings to handle various text file formats
        encodings = ['utf-8', 'utf-16', 'cp1252', 'iso-8859-1', 'ascii']
        
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    content = f.read()
                print(f"✅ Extracted text from TXT file using {encoding}: {len(content)} characters")
                return content
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # If all encodings fail, try binary mode and decode with error handling
        with open(path, 'rb') as f:
            raw_content = f.read()
            content = raw_content.decode('utf-8', errors='replace')
        
        print(f"✅ Extracted text from TXT file (with replacement chars): {len(content)} characters")
        return content
        
    except Exception as e:
        print(f"❌ Error extracting TXT text from {path}: {e}")
        return ""

def clean_email_body(email_body: str) -> str:
    """Clean and preprocess email body text"""
    if not email_body:
        return ""
    
    # Remove HTML tags if present
    import re
    clean_text = re.sub(r'<[^>]+>', '', email_body)
    
    # Remove excessive whitespace
    clean_text = re.sub(r'\n\s*\n', '\n\n', clean_text)
    clean_text = re.sub(r' +', ' ', clean_text)
    
    return clean_text.strip()

def extract_tabular_data_from_email(email_text: str) -> str:
    """Extract and format tabular data from email body"""
    # Look for common table patterns in email
    lines = email_text.split('\n')
    table_lines = []
    
    for line in lines:
        line = line.strip()
        # Check if line might be part of a table (contains multiple separators)
        if any(sep in line for sep in ['\t', '|', '  ', ':', ';']) and len(line) > 10:
            # Clean up the line for better processing
            cleaned_line = re.sub(r'\s{2,}', ' | ', line)
            table_lines.append(cleaned_line)
    
    if table_lines:
        return "=== EXTRACTED TABLE DATA ===\n" + "\n".join(table_lines)
    return ""

def process_all_content(email_body: str, attachments_info: List[Dict[str, Any]]) -> str:
    """Process email body and all attachments to create combined text"""
    content_parts = []
    
    # Process email body
    if email_body:
        cleaned_body = clean_email_body(email_body)
        if cleaned_body:
            content_parts.append("=== EMAIL BODY ===\n" + cleaned_body)
            
            # Extract any tabular data from email
            table_data = extract_tabular_data_from_email(cleaned_body)
            if table_data:
                content_parts.append(table_data)
    
    # Process attachments
    for attachment_path in attachments_info:
        if not attachment_path or not os.path.exists(attachment_path):
            continue
            
        filename = os.path.basename(attachment_path)
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        content_parts.append(f"=== ATTACHMENT: {filename} ===")
        
        if file_ext == 'pdf':
            pdf_text = extract_text_from_pdf(attachment_path)
            if pdf_text:
                content_parts.append(pdf_text)
        elif file_ext in ['xls', 'xlsx']:
            excel_text = extract_text_from_excel(attachment_path)
            if excel_text:
                content_parts.append(excel_text)
        elif file_ext == 'docx':
            docx_text = extract_text_from_docx(attachment_path)
            if docx_text:
                content_parts.append(docx_text)
        elif file_ext == 'txt':
            txt_text = extract_text_from_txt(attachment_path)
            if txt_text:
                content_parts.append(txt_text)
        else:
            print(f"⚠️  Unsupported file type: {filename}")
    
    combined_content = "\n\n".join(content_parts)
    print(f"✅ Combined content length: {len(combined_content)} characters")
    
    return combined_content